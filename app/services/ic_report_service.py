"""V4-5: Investment Committee / Credit Memo Reporting Pack service.

Pure computation — no persistence side effects, no engine changes.
Assembles report data from canonical WaterfallResult and ProjectInputs.
"""
from __future__ import annotations

import io
from datetime import date
from typing import Any, Optional

# ─── Version constants ────────────────────────────────────────────────────────

ENGINE_VERSION = "V4-5"
MODEL_VERSION = "Finco One 4.5"


# ─── Data assembly ────────────────────────────────────────────────────────────

def build_exec_summary(
    proj: Any,
    result: Any,
    scenario_name: str = "Base",
    lender_result: Optional[Any] = None,
) -> dict[str, Any]:
    """Assemble executive summary data from canonical engine output.

    Returns a dict with all fields needed for Part A template.
    No recalculation — values sourced directly from WaterfallResult.
    """
    fin = proj.financing
    tech = proj.technical
    info = proj.info
    rev = proj.revenue
    tax = proj.tax

    gearing = getattr(fin, "gearing_ratio", 0.0) or 0.0
    total_capex = proj.capex.total_capex
    debt_keur = total_capex * gearing
    equity_keur = total_capex * (1.0 - gearing)

    # Payback: first period where cum_distribution >= equity invested
    payback_years = _calc_payback(result, equity_keur)

    # Technology detection
    tech_name = _detect_technology(info)

    # Traffic-light RAG for summary
    avg_dscr = result.actual_avg_dscr
    min_dscr = result.min_dscr
    equity_irr = result.equity_irr
    rag_financial = "green" if equity_irr and equity_irr >= 0.08 else ("amber" if equity_irr and equity_irr >= 0.05 else "red")
    rag_debt = "green" if min_dscr and min_dscr >= 1.15 else ("amber" if min_dscr and min_dscr >= 1.05 else "red")
    rag_tax = "green" if result.total_tax_keur is not None else "amber"
    rag_covenants = "green" if min_dscr and min_dscr >= 1.10 else ("amber" if min_dscr and min_dscr >= 1.05 else "red")

    return {
        # Project
        "project_name": info.name,
        "company": info.company,
        "technology": tech_name,
        "country": info.country_iso,
        "cod_date": str(info.cod_date),
        "horizon_years": info.horizon_years,
        "capacity_mw": tech.capacity_mw,
        "scenario_name": scenario_name,
        "report_date": str(date.today()),
        "engine_version": ENGINE_VERSION,
        "model_version": MODEL_VERSION,
        # Capital structure
        "total_capex_keur": total_capex,
        "debt_keur": debt_keur,
        "equity_keur": equity_keur,
        "gearing_pct": gearing * 100.0,
        "equity_pct": (1.0 - gearing) * 100.0,
        "senior_tenor_years": getattr(fin, "senior_tenor_years", None),
        "corporate_tax_rate_pct": tax.corporate_rate * 100.0,
        # KPIs
        "project_irr": result.project_irr,
        "equity_irr": result.equity_irr,
        "equity_npv": result.equity_npv,
        "avg_dscr": result.actual_avg_dscr,
        "min_dscr": result.min_dscr,
        "min_llcr": result.min_llcr,
        "min_plcr": result.min_plcr,
        "total_distribution_keur": result.total_distribution_keur,
        "total_tax_keur": result.total_tax_keur,
        "total_revenue_keur": result.total_revenue_keur,
        "total_ebitda_keur": result.total_ebitda_keur,
        "total_opex_keur": getattr(result, "total_opex_keur", None),
        "total_senior_ds_keur": result.total_senior_ds_keur,
        "periods_in_lockup": result.periods_in_lockup,
        "payback_years": payback_years,
        # PPA
        "ppa_tariff": rev.ppa_base_tariff,
        "ppa_term_years": rev.ppa_term_years,
        # Traffic lights
        "rag_financial": rag_financial,
        "rag_debt": rag_debt,
        "rag_tax": rag_tax,
        "rag_covenants": rag_covenants,
    }


def build_ic_pack(
    proj: Any,
    result: Any,
    scenario_name: str = "Base",
    sensitivity_result: Optional[dict] = None,
    covenant_periods: Optional[list] = None,
) -> dict[str, Any]:
    """Assemble Investment Committee pack data.

    Sections: executive summary, scenario overview, key risks, financial KPIs,
    sensitivity summary, financial statements outline, debt summary, tax summary,
    distribution summary, key assumptions.
    """
    exec_summary = build_exec_summary(proj, result, scenario_name)

    # Debt schedule summary
    debt_summary = _build_debt_summary(result, proj)

    # Tax summary
    tax_summary = _build_tax_summary(result, proj)

    # Distribution schedule (annual aggregation from periods)
    dist_summary = _build_distribution_summary(result)

    # Key assumptions
    assumptions = _build_key_assumptions(proj)

    # Key risks
    risks = _build_key_risks(proj, result)

    # Sensitivity headline (if available)
    sens_headline = None
    if sensitivity_result:
        sens_headline = _build_sensitivity_headline(sensitivity_result)

    return {
        "exec_summary": exec_summary,
        "debt_summary": debt_summary,
        "tax_summary": tax_summary,
        "dist_summary": dist_summary,
        "assumptions": assumptions,
        "risks": risks,
        "sens_headline": sens_headline,
        "covenant_periods": covenant_periods or [],
        "scenario_name": scenario_name,
        "report_date": str(date.today()),
        "engine_version": ENGINE_VERSION,
        "model_version": MODEL_VERSION,
    }


def build_credit_pack(
    proj: Any,
    result: Any,
    scenario_name: str = "Base",
    lender_result: Optional[dict] = None,
    covenant_periods: Optional[list] = None,
    covenant_thresholds: Optional[dict] = None,
) -> dict[str, Any]:
    """Assemble Credit Committee pack (bank-oriented).

    Focus: debt sizing, DSCR/LLCR/PLCR, lockup, distribution restrictions,
    cash sweep, covenant history, stress scenarios.
    """
    from app.services.lender_case_service import (
        DSCR_EVENT_OF_DEFAULT, DSCR_LOCKUP, DSCR_DISTRIBUTION, DSCR_CASH_SWEEP,
    )

    exec_summary = build_exec_summary(proj, result, scenario_name)
    debt_summary = _build_debt_summary(result, proj)

    thresholds = covenant_thresholds or {
        "event_of_default": DSCR_EVENT_OF_DEFAULT,
        "lockup": DSCR_LOCKUP,
        "distribution": DSCR_DISTRIBUTION,
        "cash_sweep": DSCR_CASH_SWEEP,
    }

    # Covenant analytics
    cov_periods = covenant_periods or []
    breach_count = sum(1 for p in cov_periods if p.get("rag") == "breach")
    warning_count = sum(1 for p in cov_periods if p.get("rag") in ("warning", "caution"))
    lockup_count = sum(1 for p in cov_periods if p.get("lockup"))

    # Lender stress KPIs delta
    stress_delta = None
    if lender_result:
        stress_delta = _build_stress_delta(result, lender_result)

    return {
        "exec_summary": exec_summary,
        "debt_summary": debt_summary,
        "covenant_periods": cov_periods,
        "covenant_thresholds": thresholds,
        "breach_count": breach_count,
        "warning_count": warning_count,
        "lockup_count": lockup_count,
        "stress_delta": stress_delta,
        "scenario_name": scenario_name,
        "report_date": str(date.today()),
        "engine_version": ENGINE_VERSION,
        "model_version": MODEL_VERSION,
    }


# ─── Export functions ─────────────────────────────────────────────────────────

def export_report_xlsx(
    proj: Any,
    result: Any,
    scenario_name: str = "Base",
    covenant_periods: Optional[list] = None,
) -> bytes:
    """Export IC/credit report pack as Excel workbook.

    Sheets: Executive Summary, Covenant Schedule, Debt Schedule,
    Distribution Schedule, Tax Summary.
    """
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = openpyxl.Workbook()

    # Styles
    hdr_font = Font(bold=True, size=10, color="FFFFFF")
    hdr_fill = PatternFill("solid", fgColor="1E3A5F")
    sub_font = Font(bold=True, size=9, color="1E3A5F")
    sub_fill = PatternFill("solid", fgColor="D6E4F0")
    num_align = Alignment(horizontal="right")

    def _hdr_row(ws, row, cols):
        for c, label in enumerate(cols, 1):
            cell = ws.cell(row=row, column=c, value=label)
            cell.font = hdr_font
            cell.fill = hdr_fill
            cell.alignment = Alignment(horizontal="center")

    def _sub_row(ws, row, cols):
        for c, label in enumerate(cols, 1):
            cell = ws.cell(row=row, column=c, value=label)
            cell.font = sub_font
            cell.fill = sub_fill

    def _num(ws, row, col, val, fmt="#,##0.00"):
        if val is None:
            return
        cell = ws.cell(row=row, column=col, value=val)
        cell.number_format = fmt
        cell.alignment = num_align

    exec_s = build_exec_summary(proj, result, scenario_name)

    # ── Sheet 1: Executive Summary ──
    ws1 = wb.active
    ws1.title = "Executive Summary"
    ws1.column_dimensions["A"].width = 28
    ws1.column_dimensions["B"].width = 22

    title_font = Font(bold=True, size=13, color="1E3A5F")
    ws1["A1"].value = f"Finco One — {exec_s['project_name']}"
    ws1["A1"].font = title_font
    ws1["A2"].value = f"Report date: {exec_s['report_date']} | Scenario: {scenario_name}"
    ws1["A2"].font = Font(italic=True, size=9, color="6B7280")

    rows = [
        ("", ""),
        ("PROJECT", ""),
        ("Technology", exec_s["technology"]),
        ("Country", exec_s["country"]),
        ("COD", exec_s["cod_date"]),
        ("Capacity (MW)", exec_s["capacity_mw"]),
        ("Horizon (yrs)", exec_s["horizon_years"]),
        ("", ""),
        ("CAPITAL STRUCTURE", ""),
        ("Total CAPEX (kEUR)", exec_s["total_capex_keur"]),
        ("Senior Debt (kEUR)", exec_s["debt_keur"]),
        ("Equity (kEUR)", exec_s["equity_keur"]),
        ("Gearing (%)", exec_s["gearing_pct"]),
        ("Tenor (yrs)", exec_s["senior_tenor_years"]),
        ("", ""),
        ("RETURNS", ""),
        ("Project IRR", exec_s["project_irr"]),
        ("Equity IRR", exec_s["equity_irr"]),
        ("Equity NPV (kEUR)", exec_s["equity_npv"]),
        ("Payback (yrs)", exec_s["payback_years"]),
        ("Lifetime Distributions (kEUR)", exec_s["total_distribution_keur"]),
        ("", ""),
        ("DEBT METRICS", ""),
        ("Min DSCR", exec_s["min_dscr"]),
        ("Avg DSCR", exec_s["avg_dscr"]),
        ("Min LLCR", exec_s["min_llcr"]),
        ("Min PLCR", exec_s["min_plcr"]),
        ("Periods in Lock-up", exec_s["periods_in_lockup"]),
        ("", ""),
        ("TAX", ""),
        ("Corp Tax Rate (%)", exec_s["corporate_tax_rate_pct"]),
        ("Total Tax (kEUR)", exec_s["total_tax_keur"]),
    ]

    for i, (label, val) in enumerate(rows, 4):
        ws1.cell(row=i, column=1, value=label)
        if label.isupper() and label:
            ws1.cell(row=i, column=1).font = sub_font
            ws1.cell(row=i, column=1).fill = sub_fill
        if val != "" and isinstance(val, float):
            cell = ws1.cell(row=i, column=2, value=val)
            if "IRR" in label or "DSCR" in label or "LLCR" in label or "PLCR" in label or "%" in label:
                cell.number_format = "0.00%"
            else:
                cell.number_format = "#,##0.0"
            cell.alignment = num_align
        elif val != "":
            ws1.cell(row=i, column=2, value=val)

    # ── Sheet 2: Covenant Schedule ──
    ws2 = wb.create_sheet("Covenant Schedule")
    ws2.column_dimensions["A"].width = 8
    ws2.column_dimensions["B"].width = 12
    for col in ["C", "D", "E", "F", "G", "H", "I", "J", "K"]:
        ws2.column_dimensions[col].width = 14

    _hdr_row(ws2, 1, ["Period", "Date", "DSCR", "LLCR", "PLCR",
                       "Sr Bal (kEUR)", "Sr DS (kEUR)", "Lockup",
                       "Distribution (kEUR)", "Cash Sweep (kEUR)", "Status"])

    rag_fills = {
        "breach": PatternFill("solid", fgColor="FEE2E2"),
        "warning": PatternFill("solid", fgColor="FFEDD5"),
        "caution": PatternFill("solid", fgColor="FEF9C3"),
        "ok": PatternFill("solid", fgColor="F0FDF4"),
    }

    for r_idx, p in enumerate(covenant_periods or [], 2):
        ws2.cell(r_idx, 1, p.get("period"))
        ws2.cell(r_idx, 2, p.get("date", "")[:7])
        _num(ws2, r_idx, 3, p.get("dscr"), "0.0000")
        _num(ws2, r_idx, 4, p.get("llcr"), "0.0000")
        plcr = p.get("plcr")
        _num(ws2, r_idx, 5, plcr if plcr and plcr > 0 else None, "0.0000")
        _num(ws2, r_idx, 6, p.get("senior_balance_keur"), "#,##0.0")
        _num(ws2, r_idx, 7, p.get("senior_ds_keur"), "#,##0.0")
        ws2.cell(r_idx, 8, "LOCKED" if p.get("lockup") else "")
        _num(ws2, r_idx, 9, p.get("distribution_keur") or None, "#,##0.0")
        _num(ws2, r_idx, 10, p.get("cash_sweep_keur") or None, "#,##0.0")
        ws2.cell(r_idx, 11, p.get("rag", "").upper())
        rag = p.get("rag", "ok")
        fill = rag_fills.get(rag, rag_fills["ok"])
        for col in range(1, 12):
            ws2.cell(r_idx, col).fill = fill

    # ── Sheet 3: Distribution Schedule ──
    ws3 = wb.create_sheet("Distribution Schedule")
    ws3.column_dimensions["A"].width = 10
    ws3.column_dimensions["B"].width = 16
    ws3.column_dimensions["C"].width = 18
    _hdr_row(ws3, 1, ["Year", "Distribution (kEUR)", "Cumulative (kEUR)"])

    annual = _build_annual_distributions(result)
    cum = 0.0
    for r_idx, (yr, dist) in enumerate(annual, 2):
        cum += dist
        ws3.cell(r_idx, 1, yr)
        _num(ws3, r_idx, 2, dist, "#,##0.0")
        _num(ws3, r_idx, 3, cum, "#,##0.0")

    # ── Sheet 4: Debt Schedule ──
    ws4 = wb.create_sheet("Debt Schedule")
    ws4.column_dimensions["A"].width = 8
    ws4.column_dimensions["B"].width = 12
    for col in ["C", "D", "E", "F"]:
        ws4.column_dimensions[col].width = 18
    _hdr_row(ws4, 1, ["Period", "Date", "Opening Balance (kEUR)",
                       "Interest (kEUR)", "Principal (kEUR)", "Closing Balance (kEUR)"])

    prev_bal = None
    op_periods = [p for p in result.periods if p.is_operation]
    for r_idx, p in enumerate(op_periods, 2):
        ws4.cell(r_idx, 1, p.period)
        ws4.cell(r_idx, 2, str(p.date)[:7] if p.date else "")
        open_bal = prev_bal if prev_bal is not None else getattr(p, "senior_balance_keur", 0)
        _num(ws4, r_idx, 3, open_bal, "#,##0.0")
        _num(ws4, r_idx, 4, getattr(p, "senior_interest_keur", None), "#,##0.0")
        _num(ws4, r_idx, 5, getattr(p, "senior_principal_keur", None), "#,##0.0")
        close_bal = getattr(p, "senior_balance_keur", None)
        _num(ws4, r_idx, 6, close_bal, "#,##0.0")
        prev_bal = close_bal

    # ── Sheet 5: Tax Summary ──
    ws5 = wb.create_sheet("Tax Summary")
    ws5.column_dimensions["A"].width = 8
    ws5.column_dimensions["B"].width = 12
    for col in ["C", "D", "E"]:
        ws5.column_dimensions[col].width = 18
    _hdr_row(ws5, 1, ["Period", "Date", "EBITDA (kEUR)", "Tax Charge (kEUR)", "Tax Loss CF (kEUR)"])

    for r_idx, p in enumerate(op_periods, 2):
        ws5.cell(r_idx, 1, p.period)
        ws5.cell(r_idx, 2, str(p.date)[:7] if p.date else "")
        _num(ws5, r_idx, 3, getattr(p, "ebitda_keur", None), "#,##0.0")
        _num(ws5, r_idx, 4, getattr(p, "tax_keur", None), "#,##0.0")
        _num(ws5, r_idx, 5, getattr(p, "r100_carryforward_keur", None), "#,##0.0")

    # ── Sheet 6: P&L Statement ──
    from domain.financial_statements import assemble_financial_statements
    try:
        fs = assemble_financial_statements(result)
        ws6 = wb.create_sheet("P&L Statement")
        ws6.column_dimensions["A"].width = 30
        for col in ["B", "C", "D", "E", "F"]:
            ws6.column_dimensions[col].width = 16
        pnl_periods = list(fs.pnl.periods)
        _hdr_row(ws6, 1, ["Metric"] + [str(p.date)[:7] for p in pnl_periods[:20]])
        pnl_rows = [
            ("Revenue (kEUR)",          [p.revenues_keur for p in pnl_periods]),
            ("OPEX (kEUR)",             [p.operating_expenses_keur for p in pnl_periods]),
            ("Depreciation (kEUR)",     [p.depreciation_keur for p in pnl_periods]),
            ("EBIT (kEUR)",             [p.ebit_keur for p in pnl_periods]),
            ("Senior Interest (kEUR)",  [p.senior_interest_expense_keur for p in pnl_periods]),
            ("SHL Interest (kEUR)",     [p.shl_interest_expense_keur for p in pnl_periods]),
            ("EBT (kEUR)",              [p.earnings_before_tax_keur for p in pnl_periods]),
            ("Taxable Income (kEUR)",   [p.taxable_income_before_losses_keur for p in pnl_periods]),
            ("Net Income (kEUR)",       [p.net_income_keur for p in pnl_periods]),
            ("Net Dividends (kEUR)",    [p.net_dividends_keur for p in pnl_periods]),
        ]
        for r_idx, (label, vals) in enumerate(pnl_rows, 2):
            ws6.cell(r_idx, 1, label)
            for c_idx, v in enumerate(vals[:20], 2):
                _num(ws6, r_idx, c_idx, v)

        # ── Sheet 7: Balance Sheet ──
        ws7 = wb.create_sheet("Balance Sheet")
        ws7.column_dimensions["A"].width = 30
        for col in ["B", "C", "D", "E", "F"]:
            ws7.column_dimensions[col].width = 16
        bs_periods = list(fs.balance_sheet.periods)
        _hdr_row(ws7, 1, ["Metric"] + [str(p.date)[:7] for p in bs_periods[:20]])
        bs_rows = [
            ("Gross Fixed Assets (kEUR)",        [p.gross_fixed_assets_keur for p in bs_periods]),
            ("Accumulated Depreciation (kEUR)",  [p.accumulated_depreciation_keur for p in bs_periods]),
            ("Net Fixed Assets (kEUR)",          [p.net_fixed_assets_keur for p in bs_periods]),
            ("Cash (kEUR)",                      [p.cash_keur for p in bs_periods]),
            ("Total Assets (kEUR)",              [p.total_assets_keur for p in bs_periods]),
            ("Senior Debt Balance (kEUR)",       [p.senior_balance_keur for p in bs_periods]),
            ("SHL Balance (kEUR)",               [p.shl_balance_keur for p in bs_periods]),
            ("Retained Earnings (kEUR)",         [p.retained_earnings_keur for p in bs_periods]),
            ("Total Liabilities + Equity (kEUR)",[p.total_liabilities_equity_keur for p in bs_periods]),
            ("Balance Check (kEUR)",             [p.balance_check_keur for p in bs_periods]),
        ]
        for r_idx, (label, vals) in enumerate(bs_rows, 2):
            ws7.cell(r_idx, 1, label)
            for c_idx, v in enumerate(vals[:20], 2):
                _num(ws7, r_idx, c_idx, v)

        # ── Sheet 8: Cash Flow Statement ──
        ws8 = wb.create_sheet("Cash Flow Statement")
        ws8.column_dimensions["A"].width = 30
        for col in ["B", "C", "D", "E", "F"]:
            ws8.column_dimensions[col].width = 16
        cf_periods = list(fs.pf_cash_waterfall.periods)
        _hdr_row(ws8, 1, ["Metric"] + [str(p.date)[:7] for p in cf_periods[:20]])
        cf_rows = [
            ("Revenue Cash (kEUR)",       [p.revenue_cash_keur for p in cf_periods]),
            ("OPEX Cash (kEUR)",          [p.opex_cash_keur for p in cf_periods]),
            ("EBITDA Cash (kEUR)",        [p.ebitda_cash_keur for p in cf_periods]),
            ("Cash Tax (kEUR)",           [p.cash_tax_keur for p in cf_periods]),
            ("FCF Banks (kEUR)",          [p.fcf_banks_keur for p in cf_periods]),
            ("Senior Debt Service (kEUR)",[p.senior_total_ds_keur for p in cf_periods]),
            ("FCF for SHL (kEUR)",        [p.fcf_for_shl_keur for p in cf_periods]),
            ("SHL Cash Interest (kEUR)",  [p.shl_cash_interest_keur for p in cf_periods]),
            ("SHL Principal (kEUR)",      [p.shl_principal_keur for p in cf_periods]),
            ("Net Dividends (kEUR)",      [p.net_dividends_keur for p in cf_periods]),
        ]
        for r_idx, (label, vals) in enumerate(cf_rows, 2):
            ws8.cell(r_idx, 1, label)
            for c_idx, v in enumerate(vals[:20], 2):
                _num(ws8, r_idx, c_idx, v)
    except Exception:
        pass  # FS assembly failure must not break the existing sheets

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def export_report_docx(
    proj: Any,
    result: Any,
    scenario_name: str = "Base",
    report_type: str = "ic",
    covenant_periods: Optional[list] = None,
) -> bytes:
    """Export report as DOCX. report_type: 'ic' or 'credit'."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    import docx.oxml as oxml

    exec_s = build_exec_summary(proj, result, scenario_name)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title
    title_label = "Investment Committee Pack" if report_type == "ic" else "Credit Committee Pack"
    title = doc.add_heading(f"Finco One — {title_label}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(f"{exec_s['project_name']} | {scenario_name} | {exec_s['report_date']}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph()

    def _section_heading(text):
        h = doc.add_heading(text, level=1)

    def _kv(label, value):
        p = doc.add_paragraph()
        run_l = p.add_run(f"{label}: ")
        run_l.bold = True
        run_l.font.size = Pt(10)
        run_v = p.add_run(str(value) if value is not None else "—")
        run_v.font.size = Pt(10)

    def _fmt_pct(v):
        return f"{v*100:.2f}%" if v is not None else "—"

    def _fmt_keur(v):
        return f"{v:,.0f} kEUR" if v is not None else "—"

    def _fmt_x(v):
        return f"{v:.4f}x" if v is not None else "—"

    # Section 1 — Project Overview
    _section_heading("1. Project Overview")
    _kv("Project", exec_s["project_name"])
    _kv("Company", exec_s["company"])
    _kv("Technology", exec_s["technology"])
    _kv("Country", exec_s["country"])
    _kv("COD", exec_s["cod_date"])
    _kv("Capacity", f"{exec_s['capacity_mw']:.1f} MW")
    _kv("Horizon", f"{exec_s['horizon_years']} years")
    _kv("PPA Tariff", f"{exec_s['ppa_tariff']:.2f} EUR/MWh")
    _kv("PPA Term", f"{exec_s['ppa_term_years']} years")

    # Section 2 — Capital Structure
    _section_heading("2. Capital Structure")
    _kv("Total CAPEX", _fmt_keur(exec_s["total_capex_keur"]))
    _kv("Senior Debt", _fmt_keur(exec_s["debt_keur"]))
    _kv("Equity", _fmt_keur(exec_s["equity_keur"]))
    _kv("Gearing", f"{exec_s['gearing_pct']:.1f}%")
    _kv("Senior Tenor", f"{exec_s['senior_tenor_years']} years" if exec_s["senior_tenor_years"] else "—")
    _kv("Corporate Tax Rate", f"{exec_s['corporate_tax_rate_pct']:.1f}%")

    # Section 3 — Financial Returns
    _section_heading("3. Financial Returns")
    _kv("Project IRR", _fmt_pct(exec_s["project_irr"]))
    _kv("Equity IRR", _fmt_pct(exec_s["equity_irr"]))
    _kv("Equity NPV", _fmt_keur(exec_s["equity_npv"]))
    _kv("Payback", f"{exec_s['payback_years']:.1f} years" if exec_s["payback_years"] else "—")
    _kv("Lifetime Distributions", _fmt_keur(exec_s["total_distribution_keur"]))
    _kv("Total Tax", _fmt_keur(exec_s["total_tax_keur"]))

    # Section 4 — Debt Metrics (credit pack: more detail)
    _section_heading("4. Debt Metrics")
    _kv("Min DSCR", _fmt_x(exec_s["min_dscr"]))
    _kv("Avg DSCR", _fmt_x(exec_s["avg_dscr"]))
    _kv("Min LLCR", _fmt_x(exec_s["min_llcr"]))
    _kv("Min PLCR", _fmt_x(exec_s["min_plcr"]))
    _kv("Periods in Lock-up", str(exec_s["periods_in_lockup"]))

    if report_type == "credit" and covenant_periods:
        _section_heading("5. Covenant Schedule Summary")
        from app.services.lender_case_service import DSCR_EVENT_OF_DEFAULT, DSCR_LOCKUP, DSCR_DISTRIBUTION
        breach_count = sum(1 for p in covenant_periods if p.get("rag") == "breach")
        warn_count = sum(1 for p in covenant_periods if p.get("rag") in ("warning", "caution"))
        lock_count = sum(1 for p in covenant_periods if p.get("lockup"))
        _kv("Total Operating Periods", str(len(covenant_periods)))
        _kv("Breach Periods (DSCR < 1.05x)", str(breach_count))
        _kv("Warning/Caution Periods", str(warn_count))
        _kv("Lock-up Periods", str(lock_count))

        # Covenant table (first 20 rows)
        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, h in enumerate(["Period", "Date", "DSCR", "LLCR", "Distribution (kEUR)", "Status"]):
            hdr[i].text = h
            for run in hdr[i].paragraphs[0].runs:
                run.bold = True

        for p in covenant_periods[:20]:
            row = table.add_row().cells
            row[0].text = str(p.get("period", ""))
            row[1].text = p.get("date", "")[:7]
            row[2].text = f"{p['dscr']:.4f}x" if p.get("dscr") is not None else "—"
            row[3].text = f"{p['llcr']:.4f}x" if p.get("llcr") is not None else "—"
            dist = p.get("distribution_keur", 0)
            row[4].text = f"{dist:,.1f}" if dist and dist > 0 else "—"
            row[5].text = p.get("rag", "").upper()

    # Footer metadata
    doc.add_paragraph()
    footer_p = doc.add_paragraph(
        f"Generated by Finco One {MODEL_VERSION} | Engine: {ENGINE_VERSION} | {exec_s['report_date']}"
    )
    for run in footer_p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─── Private helpers ─────────────────────────────────────────────────────────

def _detect_technology(info: Any) -> str:
    name = (info.name or "").lower()
    if "wind" in name:
        return "Wind"
    if "solar" in name or "pv" in name:
        return "Solar PV"
    if "bess" in name or "battery" in name:
        return "BESS"
    return "Renewable"


def _calc_payback(result: Any, equity_keur: float) -> Optional[float]:
    """Return years from COD to payback (cum_distribution >= equity)."""
    if not equity_keur or equity_keur <= 0:
        return None
    op_periods = [p for p in result.periods if p.is_operation]
    if not op_periods:
        return None
    first_op = op_periods[0]
    for p in op_periods:
        cum = getattr(p, "cum_distribution_keur", 0.0) or 0.0
        if cum >= equity_keur:
            # Years from first operation period
            return (p.period - first_op.period) * 0.5  # semiannual periods
    return None


def _build_debt_summary(result: Any, proj: Any) -> dict[str, Any]:
    fin = proj.financing
    gearing = getattr(fin, "gearing_ratio", 0.0) or 0.0
    total_capex = proj.capex.total_capex
    return {
        "debt_keur": total_capex * gearing,
        "senior_tenor_years": getattr(fin, "senior_tenor_years", None),
        "total_senior_ds_keur": result.total_senior_ds_keur,
        "min_dscr": result.min_dscr,
        "avg_dscr": result.actual_avg_dscr,
        "min_llcr": result.min_llcr,
        "min_plcr": result.min_plcr,
        "periods_in_lockup": result.periods_in_lockup,
    }


def _build_tax_summary(result: Any, proj: Any) -> dict[str, Any]:
    tax = proj.tax
    return {
        "corporate_tax_rate_pct": tax.corporate_rate * 100.0,
        "total_tax_keur": result.total_tax_keur,
        "total_ebitda_keur": result.total_ebitda_keur,
        "total_revenue_keur": result.total_revenue_keur,
    }


def _build_distribution_summary(result: Any) -> dict[str, Any]:
    annual = _build_annual_distributions(result)
    return {
        "total_keur": result.total_distribution_keur,
        "annual": annual,
        "periods_count": sum(1 for p in result.periods if p.is_operation and getattr(p, "distribution_keur", 0) > 0),
    }


def _build_annual_distributions(result: Any) -> list[tuple[int, float]]:
    """Aggregate semiannual distributions to annual buckets keyed by year_index."""
    annual: dict[int, float] = {}
    for p in result.periods:
        if not p.is_operation:
            continue
        dist = getattr(p, "distribution_keur", 0.0) or 0.0
        if dist <= 0:
            continue
        yr = getattr(p, "year_index", None)
        if yr is None:
            continue
        annual[yr] = annual.get(yr, 0.0) + dist
    return sorted(annual.items())


def _build_key_assumptions(proj: Any) -> list[dict[str, str]]:
    tech = proj.technical
    rev = proj.revenue
    fin = proj.financing
    rows = [
        ("P50 Operating Hours", f"{tech.operating_hours_p50:,.0f} hrs/yr"),
        ("Plant Availability", f"{tech.plant_availability*100:.1f}%"),
        ("PPA Tariff", f"{rev.ppa_base_tariff:.2f} EUR/MWh"),
        ("PPA Term", f"{rev.ppa_term_years} years"),
        ("Gearing", f"{getattr(fin, 'gearing_ratio', 0)*100:.1f}%"),
        ("Senior Tenor", f"{getattr(fin, 'senior_tenor_years', '?')} years"),
    ]
    return [{"label": l, "value": v} for l, v in rows]


def _build_key_risks(proj: Any, result: Any) -> list[dict[str, str]]:
    risks = []
    if result.min_dscr and result.min_dscr < 1.20:
        risks.append({"risk": "Debt service headroom", "severity": "medium",
                      "note": f"Min DSCR {result.min_dscr:.2f}x is below 1.20x buffer"})
    if result.periods_in_lockup and result.periods_in_lockup > 0:
        risks.append({"risk": "Lockup periods", "severity": "medium",
                      "note": f"{result.periods_in_lockup} periods with restricted distributions"})
    if result.equity_irr and result.equity_irr < 0.08:
        risks.append({"risk": "Equity return below threshold", "severity": "high",
                      "note": f"Equity IRR {result.equity_irr*100:.2f}% below 8% hurdle"})
    rev = proj.revenue
    if rev.ppa_term_years and getattr(proj.info, "horizon_years", 25) > rev.ppa_term_years:
        risks.append({"risk": "Merchant tail exposure", "severity": "medium",
                      "note": f"PPA term {rev.ppa_term_years}yr < horizon {proj.info.horizon_years}yr"})
    if not risks:
        risks.append({"risk": "No material risks flagged", "severity": "low", "note": "Base case within covenants"})
    return risks


def _build_sensitivity_headline(sensitivity_result: dict) -> list[dict[str, Any]]:
    """Top 5 sensitivity rows by equity IRR impact."""
    rows = sensitivity_result.get("rows", [])
    irr_rows = []
    for row in rows:
        if row.get("error"):
            continue
        delta = row.get("deltas", {}).get("equity_irr")
        if delta is not None:
            irr_rows.append({"label": row["shock_label"], "level": row["level_pct"], "delta": delta})
    irr_rows.sort(key=lambda x: abs(x["delta"]), reverse=True)
    return irr_rows[:5]


def _build_stress_delta(base_result: Any, lender_result: dict) -> dict[str, Any]:
    lk = lender_result.get("kpis", {})
    return {
        "equity_irr_base": base_result.equity_irr,
        "equity_irr_stress": lk.get("equity_irr"),
        "min_dscr_base": base_result.min_dscr,
        "min_dscr_stress": lk.get("min_dscr"),
        "min_llcr_base": base_result.min_llcr,
        "min_llcr_stress": lk.get("min_llcr"),
        "adjustment_summary": lender_result.get("adjustment_summary", []),
    }
