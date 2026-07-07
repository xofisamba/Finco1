"""V4-5: Investment Committee / Credit Memo Reporting Pack tests.

Validates:
- build_exec_summary: all fields present, RAG logic, payback calc
- build_ic_pack: sections present, golden parity
- build_credit_pack: covenant integration, stress delta
- export_report_xlsx: valid workbook, expected sheets
- export_report_docx: valid DOCX bytes
- Golden-parity: waterfall core SHA unchanged
"""
from __future__ import annotations

import hashlib
import pytest


WATERFALL_CORE_SHA = "e6097569eecf4f173061bc3d8370b7b2515847c5ab0d9b6093d6401119159ef5"


# ─── Fixtures ────────────────────────────────────────────────────────────────

@pytest.fixture(scope="module")
def tuho_proj():
    from app.project_factories import create_default_tuho_wind1
    return create_default_tuho_wind1()


@pytest.fixture(scope="module")
def tuho_result(tuho_proj):
    from app.ui_runner import _build_period_engine
    from app.waterfall_runner import WaterfallRunner, WaterfallRunConfig
    eng = _build_period_engine(tuho_proj)
    return WaterfallRunner(tuho_proj, eng).run(WaterfallRunConfig.from_inputs(tuho_proj, eng))


@pytest.fixture(scope="module")
def cov_periods(tuho_result):
    from app.services.lender_case_service import build_covenant_periods
    return build_covenant_periods(tuho_result)


@pytest.fixture(scope="module")
def exec_summary(tuho_proj, tuho_result):
    from app.services.ic_report_service import build_exec_summary
    return build_exec_summary(tuho_proj, tuho_result, "Base")


@pytest.fixture(scope="module")
def ic_pack(tuho_proj, tuho_result, cov_periods):
    from app.services.ic_report_service import build_ic_pack
    return build_ic_pack(tuho_proj, tuho_result, "Base", covenant_periods=cov_periods)


@pytest.fixture(scope="module")
def credit_pack(tuho_proj, tuho_result, cov_periods):
    from app.services.ic_report_service import build_credit_pack
    from app.services.lender_case_service import (
        DSCR_EVENT_OF_DEFAULT, DSCR_LOCKUP, DSCR_DISTRIBUTION, DSCR_CASH_SWEEP,
    )
    thresholds = {
        "event_of_default": DSCR_EVENT_OF_DEFAULT,
        "lockup": DSCR_LOCKUP,
        "distribution": DSCR_DISTRIBUTION,
        "cash_sweep": DSCR_CASH_SWEEP,
    }
    return build_credit_pack(tuho_proj, tuho_result, "Base",
                              covenant_periods=cov_periods, covenant_thresholds=thresholds)


# ─── Phase 51F parity guard ───────────────────────────────────────────────────

class TestPhase51FGuard:
    def test_waterfall_core_sha_unchanged(self):
        with open("app/waterfall_core.py", "rb") as f:
            sha = hashlib.sha256(f.read()).hexdigest()
        assert sha == WATERFALL_CORE_SHA, f"waterfall_core.py SHA changed: {sha}"

    def test_golden_equity_irr(self, tuho_result):
        assert abs(tuho_result.equity_irr - 0.1132) < 0.001

    def test_golden_avg_dscr(self, tuho_result):
        assert abs(tuho_result.actual_avg_dscr - 1.3786) < 0.01

    def test_golden_distributions(self, tuho_result):
        assert abs(tuho_result.total_distribution_keur - 165471) < 1000


# ─── build_exec_summary ──────────────────────────────────────────────────────

class TestBuildExecSummary:
    def test_returns_dict(self, exec_summary):
        assert isinstance(exec_summary, dict)

    def test_required_fields(self, exec_summary):
        required = [
            "project_name", "company", "technology", "country", "cod_date",
            "horizon_years", "capacity_mw", "scenario_name", "report_date",
            "engine_version", "model_version",
            "total_capex_keur", "debt_keur", "equity_keur", "gearing_pct", "equity_pct",
            "senior_tenor_years", "corporate_tax_rate_pct",
            "project_irr", "equity_irr", "equity_npv",
            "avg_dscr", "min_dscr", "min_llcr", "min_plcr",
            "total_distribution_keur", "total_tax_keur", "total_revenue_keur",
            "total_senior_ds_keur", "periods_in_lockup",
            "ppa_tariff", "ppa_term_years",
            "rag_financial", "rag_debt", "rag_tax", "rag_covenants",
        ]
        for f in required:
            assert f in exec_summary, f"Missing field: {f}"

    def test_gearing_and_equity_sum_to_100(self, exec_summary):
        assert abs(exec_summary["gearing_pct"] + exec_summary["equity_pct"] - 100.0) < 0.001

    def test_debt_equity_sum_to_capex(self, exec_summary):
        assert abs(exec_summary["debt_keur"] + exec_summary["equity_keur"]
                   - exec_summary["total_capex_keur"]) < 1.0

    def test_rag_values_valid(self, exec_summary):
        valid = {"green", "amber", "red"}
        for key in ["rag_financial", "rag_debt", "rag_tax", "rag_covenants"]:
            assert exec_summary[key] in valid, f"{key} has invalid RAG: {exec_summary[key]}"

    def test_equity_irr_matches_canonical(self, exec_summary):
        assert abs(exec_summary["equity_irr"] - 0.1132) < 0.001

    def test_rag_financial_green_for_healthy_irr(self, exec_summary):
        # TUHO IRR ~11.3% >> 8% hurdle → green
        assert exec_summary["rag_financial"] == "green"

    def test_rag_debt_green_for_healthy_dscr(self, exec_summary):
        # TUHO min DSCR ~1.38 >> 1.15 → green
        assert exec_summary["rag_debt"] == "green"

    def test_technology_detected(self, exec_summary):
        assert exec_summary["technology"] in ("Wind", "Solar PV", "BESS", "Renewable")

    def test_capacity_mw_positive(self, exec_summary):
        assert exec_summary["capacity_mw"] > 0

    def test_payback_years_positive_or_none(self, exec_summary):
        pb = exec_summary.get("payback_years")
        if pb is not None:
            assert pb > 0

    def test_engine_version_present(self, exec_summary):
        assert exec_summary["engine_version"] != ""

    def test_report_date_present(self, exec_summary):
        assert len(exec_summary["report_date"]) == 10  # YYYY-MM-DD


# ─── build_ic_pack ───────────────────────────────────────────────────────────

class TestBuildICPack:
    def test_returns_dict(self, ic_pack):
        assert isinstance(ic_pack, dict)

    def test_required_top_level_keys(self, ic_pack):
        for key in ["exec_summary", "debt_summary", "tax_summary", "dist_summary",
                    "assumptions", "risks", "covenant_periods",
                    "scenario_name", "report_date", "engine_version", "model_version"]:
            assert key in ic_pack, f"Missing key: {key}"

    def test_exec_summary_embedded(self, ic_pack):
        assert "equity_irr" in ic_pack["exec_summary"]

    def test_risks_non_empty(self, ic_pack):
        assert len(ic_pack["risks"]) > 0

    def test_risk_fields(self, ic_pack):
        for r in ic_pack["risks"]:
            assert "risk" in r and "severity" in r and "note" in r

    def test_assumptions_non_empty(self, ic_pack):
        assert len(ic_pack["assumptions"]) > 0

    def test_dist_summary_total_positive(self, ic_pack):
        assert ic_pack["dist_summary"]["total_keur"] > 0

    def test_dist_summary_annual_non_empty(self, ic_pack):
        assert len(ic_pack["dist_summary"]["annual"]) > 0

    def test_annual_distributions_positive(self, ic_pack):
        for yr, dist in ic_pack["dist_summary"]["annual"]:
            assert dist > 0

    def test_debt_summary_fields(self, ic_pack):
        ds = ic_pack["debt_summary"]
        for key in ["debt_keur", "total_senior_ds_keur", "min_dscr", "avg_dscr"]:
            assert key in ds

    def test_tax_summary_fields(self, ic_pack):
        ts = ic_pack["tax_summary"]
        assert "corporate_tax_rate_pct" in ts
        assert "total_tax_keur" in ts

    def test_no_sens_headline_without_data(self, ic_pack):
        # Built without sensitivity_result, so no sens_headline
        assert ic_pack["sens_headline"] is None

    def test_covenant_periods_present(self, ic_pack):
        assert len(ic_pack["covenant_periods"]) > 0


# ─── build_credit_pack ───────────────────────────────────────────────────────

class TestBuildCreditPack:
    def test_returns_dict(self, credit_pack):
        assert isinstance(credit_pack, dict)

    def test_required_keys(self, credit_pack):
        for key in ["exec_summary", "debt_summary", "covenant_periods",
                    "covenant_thresholds", "breach_count", "warning_count",
                    "lockup_count", "scenario_name", "report_date"]:
            assert key in credit_pack, f"Missing key: {key}"

    def test_thresholds_correct(self, credit_pack):
        th = credit_pack["covenant_thresholds"]
        assert abs(th["event_of_default"] - 1.05) < 0.001
        assert abs(th["lockup"] - 1.10) < 0.001
        assert abs(th["distribution"] - 1.15) < 0.001

    def test_no_breaches_at_base(self, credit_pack):
        assert credit_pack["breach_count"] == 0

    def test_covenant_periods_populated(self, credit_pack):
        assert len(credit_pack["covenant_periods"]) > 0

    def test_stress_delta_none_without_lender(self, credit_pack):
        assert credit_pack["stress_delta"] is None

    def test_credit_pack_with_stress(self, tuho_proj, tuho_result, cov_periods):
        from app.services.ic_report_service import build_credit_pack
        from app.services.lender_case_service import run_lender_case, DSCR_EVENT_OF_DEFAULT, DSCR_LOCKUP, DSCR_DISTRIBUTION, DSCR_CASH_SWEEP
        lender_result = run_lender_case(tuho_proj, {"ppa_haircut": 10.0})
        cp = build_credit_pack(
            tuho_proj, tuho_result, "Base",
            lender_result=lender_result,
            covenant_periods=cov_periods,
            covenant_thresholds={
                "event_of_default": DSCR_EVENT_OF_DEFAULT,
                "lockup": DSCR_LOCKUP,
                "distribution": DSCR_DISTRIBUTION,
                "cash_sweep": DSCR_CASH_SWEEP,
            },
        )
        assert cp["stress_delta"] is not None
        sd = cp["stress_delta"]
        assert sd["equity_irr_stress"] < sd["equity_irr_base"]


# ─── export_report_xlsx ──────────────────────────────────────────────────────

class TestExportReportXlsx:
    @pytest.fixture(scope="class")
    def xlsx_bytes(self, tuho_proj, tuho_result, cov_periods):
        from app.services.ic_report_service import export_report_xlsx
        return export_report_xlsx(tuho_proj, tuho_result, "Base", covenant_periods=cov_periods)

    def test_returns_bytes(self, xlsx_bytes):
        assert isinstance(xlsx_bytes, bytes)
        assert len(xlsx_bytes) > 100

    def test_valid_xlsx(self, xlsx_bytes):
        import openpyxl, io
        wb = openpyxl.load_workbook(io.BytesIO(xlsx_bytes))
        assert wb is not None

    def test_expected_sheets(self, xlsx_bytes):
        import openpyxl, io
        wb = openpyxl.load_workbook(io.BytesIO(xlsx_bytes))
        expected = {"Executive Summary", "Covenant Schedule",
                    "Distribution Schedule", "Debt Schedule", "Tax Summary"}
        assert expected.issubset(set(wb.sheetnames))

    def test_exec_summary_has_data(self, xlsx_bytes):
        import openpyxl, io
        wb = openpyxl.load_workbook(io.BytesIO(xlsx_bytes))
        ws = wb["Executive Summary"]
        assert ws["A1"].value is not None

    def test_covenant_schedule_has_data(self, xlsx_bytes):
        import openpyxl, io
        wb = openpyxl.load_workbook(io.BytesIO(xlsx_bytes))
        ws = wb["Covenant Schedule"]
        # Header row + at least one data row
        assert ws.max_row >= 2

    def test_distribution_schedule_has_data(self, xlsx_bytes):
        import openpyxl, io
        wb = openpyxl.load_workbook(io.BytesIO(xlsx_bytes))
        ws = wb["Distribution Schedule"]
        assert ws.max_row >= 2


# ─── export_report_docx ──────────────────────────────────────────────────────

class TestExportReportDocx:
    @pytest.fixture(scope="class")
    def docx_bytes_ic(self, tuho_proj, tuho_result, cov_periods):
        from app.services.ic_report_service import export_report_docx
        return export_report_docx(tuho_proj, tuho_result, "Base",
                                   report_type="ic", covenant_periods=cov_periods)

    @pytest.fixture(scope="class")
    def docx_bytes_credit(self, tuho_proj, tuho_result, cov_periods):
        from app.services.ic_report_service import export_report_docx
        return export_report_docx(tuho_proj, tuho_result, "Base",
                                   report_type="credit", covenant_periods=cov_periods)

    def test_ic_returns_bytes(self, docx_bytes_ic):
        assert isinstance(docx_bytes_ic, bytes)
        assert len(docx_bytes_ic) > 100

    def test_credit_returns_bytes(self, docx_bytes_credit):
        assert isinstance(docx_bytes_credit, bytes)
        assert len(docx_bytes_credit) > 100

    def test_ic_valid_docx(self, docx_bytes_ic):
        from docx import Document
        import io
        doc = Document(io.BytesIO(docx_bytes_ic))
        assert len(doc.paragraphs) > 0

    def test_credit_valid_docx(self, docx_bytes_credit):
        from docx import Document
        import io
        doc = Document(io.BytesIO(docx_bytes_credit))
        assert len(doc.paragraphs) > 0

    def test_ic_contains_project_name(self, docx_bytes_ic, tuho_proj):
        from docx import Document
        import io
        doc = Document(io.BytesIO(docx_bytes_ic))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert tuho_proj.info.name in full_text

    def test_credit_contains_covenant_section(self, docx_bytes_credit):
        from docx import Document
        import io
        doc = Document(io.BytesIO(docx_bytes_credit))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Covenant" in full_text


# ─── _calc_payback helper ─────────────────────────────────────────────────────

class TestCalcPayback:
    def test_payback_positive_for_profitable_project(self, tuho_proj, tuho_result):
        from app.services.ic_report_service import _calc_payback
        fin = tuho_proj.financing
        gearing = getattr(fin, "gearing_ratio", 0.0) or 0.0
        equity_keur = tuho_proj.capex.total_capex * (1.0 - gearing)
        pb = _calc_payback(tuho_result, equity_keur)
        # TUHO distributes starting late; payback may be None if never reached
        # but if not None it must be positive
        if pb is not None:
            assert pb > 0

    def test_payback_none_for_zero_equity(self, tuho_result):
        from app.services.ic_report_service import _calc_payback
        assert _calc_payback(tuho_result, 0.0) is None

    def test_payback_none_for_very_large_equity(self, tuho_result):
        from app.services.ic_report_service import _calc_payback
        pb = _calc_payback(tuho_result, 1e12)
        assert pb is None
